from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


BASE_DIR = Path(__file__).resolve().parent
FECHA = "16/05/2026"
APP = "RH360"
CODIGO_APP = "RH3"
SISTEMA = "RRHH. Recursos Humanos"
SUBSISTEMA = "GPRH. Gestión de Personal"
FICHERO_DOCX = "RRHH_GPRH_RH3_1. Documento de alcance.docx"
FICHERO_PDF = "RRHH_GPRH_RH3_1. Documento de alcance.pdf"
AUTOR = "Eduardo Pérez Castillo"

HISTORIAS = {
    "Finalizada / Histórico": [
        "Definir nombre, sistema, subsistema y código de la aplicación.",
        "Analizar las necesidades trasladadas por Recursos Humanos.",
    ],
    "Pendiente de despliegue": [
        "Como responsable de Recursos Humanos, quiero dar de alta empleados, para permitirles acceder a la aplicación.",
    ],
    "En fase de pruebas y estrés": [
        "Como empleado, quiero consultar la agenda de compañeros, para localizar a personas de otros turnos o áreas.",
        "Como empleado, quiero marcar comunicaciones como revisadas, para controlar qué información ya he atendido.",
    ],
    "En proceso de desarrollo": [
        "Como empleado, quiero crear diarios por temática, para organizar mis anotaciones laborales.",
        "Como empleado, quiero enviar comunicaciones escritas a otros compañeros, para informarles aunque no coincidamos en el mismo turno.",
        "Como administrador, quiero configurar departamentos, áreas y turnos, para organizar correctamente la información de la empresa.",
    ],
    "Pendiente de desarrollo": [
        "Como empleado, quiero buscar compañeros por nombre, departamento o turno, para encontrar rápidamente a la persona adecuada.",
        "Como empleado, quiero registrar notas e incidencias, para dejar constancia de información importante.",
        "Como empleado, quiero consultar las comunicaciones recibidas, para revisar avisos o incidencias pendientes.",
        "Como responsable de Recursos Humanos, quiero modificar los datos de empleados, para mantener actualizada la información interna.",
        "Como responsable de Recursos Humanos, quiero dar de baja empleados, para impedir el acceso a personas que ya no pertenezcan a la organización.",
        "Como administrador, quiero gestionar roles y permisos, para controlar qué puede hacer cada tipo de usuario.",
        "Como responsable o mando intermedio, quiero consultar comunicaciones relevantes de mi área, para hacer seguimiento de incidencias importantes.",
        "Como empleado, quiero consultar el histórico de mis notas y comunicaciones, para recuperar información registrada anteriormente.",
    ],
    "Backlog del producto": [
        "Como empleado, quiero solicitar una excedencia desde la aplicación, para evitar entregar formularios en papel.",
        "Como responsable de Recursos Humanos, quiero aprobar o rechazar solicitudes de excedencia, para agilizar la tramitación.",
        "Como empleado, quiero consultar el histórico de mis solicitudes de excedencia, para conocer el estado y resultado de mis trámites.",
        "Como responsable de Recursos Humanos, quiero consultar el histórico de trámites de un empleado, para revisar sus solicitudes anteriores.",
    ],
}


def set_cell_text(cell, text, bold=False, font_size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.twips)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(w.twips) for w in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_page_field(paragraph, instruction, placeholder="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    text = OxmlElement("w:t")
    text.text = placeholder
    run._r.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def add_bottom_border(paragraph, color="BFBFBF", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def next_numbering_id(numbering, tag):
    values = [int(el.get(qn(tag))) for el in numbering.findall(f".//{tag.split(':')[0]}:{tag.split(':')[1]}", numbering.nsmap) if el.get(qn(tag))]
    return max(values, default=0) + 1


def create_num(doc, num_format, level_text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    lvl.append(fmt)
    txt = OxmlElement("w:lvlText")
    txt.set(qn("w:val"), level_text)
    lvl.append(txt)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id, ilvl=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)


def add_letter_list(doc, items):
    for idx, item in enumerate(items):
        letter = chr(ord("a") + idx)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{letter}) {item}")


def add_bullet_role(doc, role, description, num_id=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(role + ": ")
    run.bold = True
    p.add_run(description)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Heading 1", 13), ("Heading 2", 11.5), ("Heading 3", 10.5)):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        style.paragraph_format.space_before = Pt(9 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Página ")
    add_page_field(footer, "PAGE", "1")
    footer.add_run("/")
    add_page_field(footer, "NUMPAGES", "1")

    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(first_footer, "PAGE", "1")
    first_footer.add_run("/")
    add_page_field(first_footer, "NUMPAGES", "1")

    header = section.header.paragraphs[0]
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.45), WD_TAB_ALIGNMENT.RIGHT)
    header.add_run(f"{FICHERO_DOCX}\tVersión: {FECHA}")
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8.5)
    add_bottom_border(header)


def build_docx(path):
    doc = Document()
    configure_doc(doc)
    doc.core_properties.title = f"Documento de alcance - {APP}"
    doc.core_properties.author = AUTOR
    doc.core_properties.subject = "Documento de alcance académico"
    doc.core_properties.keywords = "RH360, RRHH, GPRH, RH3, documento de alcance"

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Proyecto de Desarrollo de Aplicaciones Web (PPP) / Desarrollo de Aplicaciones Web")
    run.font.name = "Arial"
    run.font.size = Pt(9)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DOCUMENTO DE ALCANCE")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    app = doc.add_paragraph()
    app.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = app.add_run(APP)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)

    for line in (SISTEMA, SUBSISTEMA):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Arial"
        r.font.size = Pt(12)

    doc.add_paragraph()
    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_width(meta, [Cm(5.2), Cm(11.2)])
    data = [("Nombre del fichero:", FICHERO_DOCX), ("Fecha de esta versión:", FECHA)]
    for row, values in zip(meta.rows, data):
        for idx, text in enumerate(values):
            set_cell_text(row.cells[idx], text, bold=idx == 0, font_size=10)
            if idx == 0:
                shade_cell(row.cells[idx], "EAF2F8")

    doc.add_page_break()

    # Historial
    add_heading(doc, "Historial de revisiones", 1)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [Cm(3.2), Cm(8.8), Cm(4.4)])
    headers = ["Fecha", "Descripción", "Autor"]
    values = [FECHA, "Creación del documento", AUTOR]
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, bold=True, font_size=10)
        shade_cell(table.rows[0].cells[idx], "EAF2F8")
    for idx, v in enumerate(values):
        set_cell_text(table.rows[1].cells[idx], v, font_size=10)
    doc.add_page_break()

    # Índice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ÍNDICE")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(13)
    entries = [
        ("1 INTRODUCCIÓN", "4"),
        ("2 REQUISITOS DE USUARIO", "4"),
        ("2.1 Necesidad del negocio", "4"),
        ("2.2 Objetivos del negocio y del proyecto", "5"),
        ("2.3 Perfil de usuarios potenciales", "5"),
        ("3 DESCRIPCIÓN DEL SISTEMA", "6"),
        ("4 CRITERIOS DE ACEPTACIÓN Y PLANIFICACIÓN", "8"),
        ("4.1 Criterios de aceptación", "8"),
    ]
    for text, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)
        p.add_run("\t" + page)
    doc.add_page_break()

    # Contenido
    add_heading(doc, "1 INTRODUCCIÓN", 1)
    add_body(
        doc,
        "El objetivo de este documento es presentar una propuesta inicial para el desarrollo de RH360, "
        "una aplicación corporativa de Recursos Humanos orientada a mejorar la comunicación interna "
        "entre empleados, el registro de anotaciones laborales, la organización de diarios temáticos y "
        "la futura gestión de trámites del empleado.",
    )
    add_body(
        doc,
        "La necesidad surge a partir de la situación trasladada por el departamento de Recursos Humanos, "
        "que ha identificado dificultades para que los empleados de distintos turnos dejen constancia de "
        "notas, incidencias u observaciones dentro del ámbito corporativo.",
    )
    add_body(
        doc,
        "El documento va dirigido a la jefatura de proyecto, analistas funcionales, responsables del área "
        "de Recursos Humanos y demás personas implicadas en la valoración inicial del proyecto.",
    )

    add_heading(doc, "2 REQUISITOS DE USUARIO", 1)
    add_heading(doc, "2.1 Necesidad del negocio", 2)
    add_body(
        doc,
        "La empresa desarrolla su actividad mediante tres turnos de trabajo: mañana, tarde y noche. "
        "Esta organización provoca que determinados empleados no coincidan con compañeros de otros "
        "turnos, aunque en ocasiones necesiten trasladarles información relevante para el desempeño de "
        "sus tareas.",
    )
    add_body(
        doc,
        "Los empleados necesitan dejar notas, incidencias u observaciones por escrito, de forma que quede "
        "constancia dentro del ámbito corporativo. No se desea utilizar el teléfono personal ni el correo "
        "personal, y el correo corporativo no se considera la solución más adecuada para este caso.",
    )
    add_body(
        doc,
        "Actualmente se utilizan pósits y anotaciones en papel. Este sistema puede provocar pérdida de "
        "información, desorden en los espacios de trabajo y conflictos con el servicio de limpieza cuando "
        "las notas físicas se retiran o se desechan.",
    )
    add_body(
        doc,
        "También se ha detectado la necesidad de disponer de diarios por temática donde los empleados "
        "puedan registrar sus propias notas de trabajo. Además, Recursos Humanos necesita controlar las "
        "altas, bajas y modificaciones de empleados dentro de la aplicación.",
    )
    add_body(
        doc,
        "Por último, los trámites de excedencia se realizan actualmente en papel y resultan poco ágiles. "
        "Aunque esta necesidad no es la más urgente, debe quedar contemplada funcionalmente para una "
        "fase posterior de RH360.",
    )
    doc.add_page_break()

    add_heading(doc, "2.2 Objetivos del negocio y del proyecto", 2)
    add_body(doc, "Los objetivos esenciales que se deberán alcanzar a corto y medio plazo son:")
    add_letter_list(
        doc,
        [
            "Facilitar una vía corporativa de comunicación escrita entre empleados de distintos turnos.",
            "Permitir que los empleados puedan decidir con qué compañeros comunicarse en cada momento.",
            "Registrar notas, incidencias y observaciones laborales para dejar constancia de la información relevante.",
            "Sustituir progresivamente el uso de pósits y anotaciones en papel por un sistema organizado.",
            "Permitir la creación y consulta de diarios por temática.",
            "Facilitar una agenda o listado de contactos internos de empleados.",
            "Permitir a Recursos Humanos gestionar el alta, baja y modificación de empleados dentro de la aplicación.",
            "Mantener un histórico de comunicaciones, notas y registros relevantes.",
            "Incorporar en una fase posterior la solicitud, revisión, aprobación o rechazo de excedencias.",
            "Permitir que cada empleado consulte el histórico de sus trámites y solicitudes.",
        ],
    )

    add_heading(doc, "2.3 Perfil de usuarios potenciales", 2)
    add_body(
        doc,
        "Dadas las necesidades indicadas y los objetivos planteados, se deberán considerar los siguientes "
        "perfiles principales de usuario:",
    )
    add_bullet_role(
        doc,
        "Empleado",
        "Usuario principal de la aplicación. Puede consultar contactos internos, crear diarios, registrar notas, enviar comunicaciones y consultar la información recibida.",
    )
    add_bullet_role(
        doc,
        "Responsable de Recursos Humanos",
        "Usuario encargado de gestionar altas, bajas y modificaciones de empleados. También podrá revisar y tramitar solicitudes relacionadas con procedimientos del empleado, como excedencias.",
    )
    add_bullet_role(
        doc,
        "Responsable o mando intermedio",
        "Usuario que, si procede según la organización, puede consultar información relacionada con su equipo, área o turno, especialmente comunicaciones o incidencias relevantes.",
    )
    add_bullet_role(
        doc,
        "Administrador del sistema",
        "Usuario encargado de configurar los parámetros generales de la aplicación, gestionar roles, permisos y mantener la estructura básica del sistema.",
    )
    doc.add_page_break()

    add_heading(doc, "3 DESCRIPCIÓN DEL SISTEMA", 1)
    add_body(
        doc,
        "RH360 se plantea como una aplicación corporativa destinada a centralizar comunicaciones internas, "
        "diarios temáticos, agenda de empleados y trámites del empleado. La aplicación se organizará en "
        "áreas funcionales que faciliten un uso claro por parte de empleados, responsables de Recursos "
        "Humanos, mandos intermedios y administradores.",
    )
    areas = [
        (
            "Área de Agenda y contactos",
            [
                "Consultar empleados registrados en la aplicación.",
                "Buscar compañeros por nombre, departamento, área o turno.",
                "Visualizar información corporativa básica de contacto interno.",
                "Facilitar la localización de empleados que trabajan en otros turnos.",
            ],
        ),
        (
            "Área de Diarios",
            [
                "Crear diarios por temática.",
                "Registrar anotaciones laborales.",
                "Consultar notas propias.",
                "Editar o archivar anotaciones.",
                "Organizar la información por fecha, tema o categoría.",
            ],
        ),
        (
            "Área de Comunicaciones",
            [
                "Crear comunicaciones dirigidas a uno o varios compañeros.",
                "Consultar comunicaciones recibidas.",
                "Responder comunicaciones cuando sea necesario.",
                "Marcar comunicaciones como revisadas.",
                "Mantener un registro histórico de comunicaciones.",
            ],
        ),
        (
            "Área de Trámites del empleado",
            [
                "Solicitar excedencias.",
                "Consultar el estado de las solicitudes.",
                "Permitir a Recursos Humanos aprobar o rechazar solicitudes.",
                "Mantener un histórico de trámites realizados por cada empleado.",
            ],
        ),
        (
            "Área de Administración",
            [
                "Dar de alta empleados.",
                "Modificar datos básicos de empleados.",
                "Dar de baja empleados.",
                "Gestionar roles y permisos.",
                "Configurar áreas, departamentos o turnos.",
                "Mantener los parámetros generales de la aplicación.",
            ],
        ),
    ]
    roman_labels = ["I", "II", "III", "IV", "V"]
    for idx, (title_area, items) in enumerate(areas):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{roman_labels[idx]}) {title_area}")
        r.bold = True
        add_letter_list(doc, items)
        if title_area == "Área de Trámites del empleado":
            add_body(
                doc,
                "Esta área puede abordarse en una fase posterior, dado que la necesidad más urgente se "
                "centra en agenda, diarios, comunicaciones y administración básica de empleados.",
            )
    doc.add_page_break()

    add_heading(doc, "4 CRITERIOS DE ACEPTACIÓN Y PLANIFICACIÓN", 1)
    add_heading(doc, "4.1 Criterios de aceptación", 2)
    add_body(
        doc,
        "Se considerará aceptable una primera fase de RH360 si quedan cubiertas las funciones esenciales "
        "de agenda, diarios, comunicaciones y administración básica de empleados. En esta primera fase "
        "no será obligatorio implantar la gestión completa de excedencias, aunque deberá quedar definida "
        "para una fase posterior.",
    )
    add_letter_list(
        doc,
        [
            "La aplicación permite el acceso de empleados previamente dados de alta por Recursos Humanos.",
            "Recursos Humanos puede crear, modificar y dar de baja empleados en la aplicación.",
            "Los empleados pueden consultar una agenda o listado de contactos internos.",
            "Los empleados pueden buscar compañeros por nombre, departamento, área o turno.",
            "Los empleados pueden crear diarios por temática.",
            "Los empleados pueden registrar notas, incidencias u observaciones laborales.",
            "Los empleados pueden enviar comunicaciones escritas a otros compañeros.",
            "Los empleados pueden consultar las comunicaciones recibidas.",
            "El sistema mantiene un registro histórico básico de notas y comunicaciones.",
            "La gestión de excedencias queda definida funcionalmente para una fase posterior, aunque no sea obligatoria en la primera entrega del sistema.",
        ],
    )

    doc.save(path)


def pstyle(name, **kwargs):
    base = getSampleStyleSheet()["Normal"]
    defaults = {"fontName": "Helvetica", "fontSize": 10, "leading": 13}
    defaults.update(kwargs)
    return ParagraphStyle(name, parent=base, **defaults)


def pdf_header_footer(canvas, doc):
    width, height = A4
    page = canvas.getPageNumber()
    canvas.setFont("Helvetica", 8)
    if page > 1:
        canvas.drawString(2.2 * cm, height - 1.35 * cm, f"{FICHERO_DOCX}")
        canvas.drawRightString(width - 2.2 * cm, height - 1.35 * cm, f"Versión: {FECHA}")
        canvas.setStrokeColor(colors.lightgrey)
        canvas.line(2.2 * cm, height - 1.45 * cm, width - 2.2 * cm, height - 1.45 * cm)
    canvas.drawCentredString(width / 2, 1.2 * cm, str(page))


def list_items(items, style):
    return ListFlowable([ListItem(Paragraph(item, style)) for item in items], bulletType="a", start="a", leftIndent=18)


def build_fallback_pdf(path):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )
    story = []
    title = pstyle("Title", alignment=TA_CENTER, fontSize=18, leading=22, textColor=colors.HexColor("#1F4E79"), spaceAfter=12)
    center = pstyle("Center", alignment=TA_CENTER, fontSize=12, leading=16, spaceAfter=6)
    h1 = pstyle("H1", fontSize=13, leading=16, textColor=colors.HexColor("#1F4E79"), spaceBefore=8, spaceAfter=6)
    h2 = pstyle("H2", fontSize=11, leading=14, textColor=colors.HexColor("#1F4E79"), spaceBefore=6, spaceAfter=4)
    normal = pstyle("Body", alignment=TA_LEFT, spaceAfter=6)

    story.append(Paragraph("Proyecto de Desarrollo de Aplicaciones Web (PPP) / Desarrollo de Aplicaciones Web", pstyle("Top", fontSize=9)))
    story.append(Spacer(1, 2.0 * cm))
    story.append(Paragraph("DOCUMENTO DE ALCANCE", title))
    story.append(Paragraph(APP, center))
    story.append(Paragraph(SISTEMA, center))
    story.append(Paragraph(SUBSISTEMA, center))
    story.append(Spacer(1, 0.8 * cm))
    table = Table([["Nombre del fichero:", FICHERO_DOCX], ["Fecha de esta versión:", FECHA]], colWidths=[5 * cm, 10.5 * cm])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF2F8")),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(table)
    story.append(PageBreak())

    story.append(Paragraph("Historial de revisiones", h1))
    rev = Table([["Fecha", "Descripción", "Autor"], [FECHA, "Creación del documento", AUTOR]], colWidths=[3.2 * cm, 8.5 * cm, 4 * cm])
    rev.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2F8")),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(rev)
    story.append(PageBreak())

    story.append(Paragraph("ÍNDICE", pstyle("IndexTitle", alignment=TA_CENTER, fontSize=13, leading=16, spaceAfter=12)))
    for text, page in [
        ("1 INTRODUCCIÓN", "4"),
        ("2 REQUISITOS DE USUARIO", "4"),
        ("2.1 Necesidad del negocio", "4"),
        ("2.2 Objetivos del negocio y del proyecto", "5"),
        ("2.3 Perfil de usuarios potenciales", "5"),
        ("3 DESCRIPCIÓN DEL SISTEMA", "6"),
        ("4 CRITERIOS DE ACEPTACIÓN Y PLANIFICACIÓN", "8"),
        ("4.1 Criterios de aceptación", "8"),
    ]:
        story.append(Paragraph(f"{text} ........................................................ {page}", normal))
    story.append(PageBreak())

    paragraphs = [
        ("1 INTRODUCCIÓN", h1),
        ("El objetivo de este documento es presentar una propuesta inicial para el desarrollo de RH360, una aplicación corporativa de Recursos Humanos orientada a mejorar la comunicación interna entre empleados, el registro de anotaciones laborales, la organización de diarios temáticos y la futura gestión de trámites del empleado.", normal),
        ("La necesidad surge a partir de la situación trasladada por el departamento de Recursos Humanos, que ha identificado dificultades para que los empleados de distintos turnos dejen constancia de notas, incidencias u observaciones dentro del ámbito corporativo.", normal),
        ("El documento va dirigido a la jefatura de proyecto, analistas funcionales, responsables del área de Recursos Humanos y demás personas implicadas en la valoración inicial del proyecto.", normal),
        ("2 REQUISITOS DE USUARIO", h1),
        ("2.1 Necesidad del negocio", h2),
        ("La empresa desarrolla su actividad mediante tres turnos de trabajo: mañana, tarde y noche. Esta organización provoca que determinados empleados no coincidan con compañeros de otros turnos, aunque en ocasiones necesiten trasladarles información relevante para el desempeño de sus tareas.", normal),
        ("Los empleados necesitan dejar notas, incidencias u observaciones por escrito, de forma que quede constancia dentro del ámbito corporativo. No se desea utilizar el teléfono personal ni el correo personal, y el correo corporativo no se considera la solución más adecuada para este caso.", normal),
        ("Actualmente se utilizan pósits y anotaciones en papel. Este sistema puede provocar pérdida de información, desorden en los espacios de trabajo y conflictos con el servicio de limpieza cuando las notas físicas se retiran o se desechan.", normal),
        ("También se ha detectado la necesidad de disponer de diarios por temática donde los empleados puedan registrar sus propias notas de trabajo. Además, Recursos Humanos necesita controlar las altas, bajas y modificaciones de empleados dentro de la aplicación.", normal),
        ("Por último, los trámites de excedencia se realizan actualmente en papel y resultan poco ágiles. Aunque esta necesidad no es la más urgente, debe quedar contemplada funcionalmente para una fase posterior de RH360.", normal),
    ]
    for text, style in paragraphs:
        story.append(Paragraph(text, style))
    story.append(PageBreak())

    story.append(Paragraph("2.2 Objetivos del negocio y del proyecto", h2))
    story.append(Paragraph("Los objetivos esenciales que se deberán alcanzar a corto y medio plazo son:", normal))
    story.append(list_items([
        "Facilitar una vía corporativa de comunicación escrita entre empleados de distintos turnos.",
        "Permitir que los empleados puedan decidir con qué compañeros comunicarse en cada momento.",
        "Registrar notas, incidencias y observaciones laborales para dejar constancia de la información relevante.",
        "Sustituir progresivamente el uso de pósits y anotaciones en papel por un sistema organizado.",
        "Permitir la creación y consulta de diarios por temática.",
        "Facilitar una agenda o listado de contactos internos de empleados.",
        "Permitir a Recursos Humanos gestionar el alta, baja y modificación de empleados dentro de la aplicación.",
        "Mantener un histórico de comunicaciones, notas y registros relevantes.",
        "Incorporar en una fase posterior la solicitud, revisión, aprobación o rechazo de excedencias.",
        "Permitir que cada empleado consulte el histórico de sus trámites y solicitudes.",
    ], normal))
    story.append(Paragraph("2.3 Perfil de usuarios potenciales", h2))
    story.append(Paragraph("Dadas las necesidades indicadas y los objetivos planteados, se deberán considerar los siguientes perfiles principales de usuario:", normal))
    story.append(Paragraph("<b>Empleado:</b> Usuario principal de la aplicación. Puede consultar contactos internos, crear diarios, registrar notas, enviar comunicaciones y consultar la información recibida.", normal))
    story.append(Paragraph("<b>Responsable de Recursos Humanos:</b> Usuario encargado de gestionar altas, bajas y modificaciones de empleados. También podrá revisar y tramitar solicitudes relacionadas con procedimientos del empleado, como excedencias.", normal))
    story.append(Paragraph("<b>Responsable o mando intermedio:</b> Usuario que, si procede según la organización, puede consultar información relacionada con su equipo, área o turno, especialmente comunicaciones o incidencias relevantes.", normal))
    story.append(Paragraph("<b>Administrador del sistema:</b> Usuario encargado de configurar los parámetros generales de la aplicación, gestionar roles, permisos y mantener la estructura básica del sistema.", normal))
    story.append(PageBreak())

    story.append(Paragraph("3 DESCRIPCIÓN DEL SISTEMA", h1))
    story.append(Paragraph("RH360 se plantea como una aplicación corporativa destinada a centralizar comunicaciones internas, diarios temáticos, agenda de empleados y trámites del empleado. La aplicación se organizará en áreas funcionales que faciliten un uso claro por parte de empleados, responsables de Recursos Humanos, mandos intermedios y administradores.", normal))
    for idx, (title_area, items) in enumerate([
        ("I) Área de Agenda y contactos", ["Consultar empleados registrados en la aplicación.", "Buscar compañeros por nombre, departamento, área o turno.", "Visualizar información corporativa básica de contacto interno.", "Facilitar la localización de empleados que trabajan en otros turnos."]),
        ("II) Área de Diarios", ["Crear diarios por temática.", "Registrar anotaciones laborales.", "Consultar notas propias.", "Editar o archivar anotaciones.", "Organizar la información por fecha, tema o categoría."]),
        ("III) Área de Comunicaciones", ["Crear comunicaciones dirigidas a uno o varios compañeros.", "Consultar comunicaciones recibidas.", "Responder comunicaciones cuando sea necesario.", "Marcar comunicaciones como revisadas.", "Mantener un registro histórico de comunicaciones."]),
        ("IV) Área de Trámites del empleado", ["Solicitar excedencias.", "Consultar el estado de las solicitudes.", "Permitir a Recursos Humanos aprobar o rechazar solicitudes.", "Mantener un histórico de trámites realizados por cada empleado."]),
        ("V) Área de Administración", ["Dar de alta empleados.", "Modificar datos básicos de empleados.", "Dar de baja empleados.", "Gestionar roles y permisos.", "Configurar áreas, departamentos o turnos.", "Mantener los parámetros generales de la aplicación."]),
    ]):
        story.append(Paragraph(f"<b>{title_area}</b>", normal))
        story.append(list_items(items, normal))
        if idx == 3:
            story.append(Paragraph("Esta área puede abordarse en una fase posterior, dado que la necesidad más urgente se centra en agenda, diarios, comunicaciones y administración básica de empleados.", normal))
    story.append(PageBreak())

    story.append(Paragraph("4 CRITERIOS DE ACEPTACIÓN Y PLANIFICACIÓN", h1))
    story.append(Paragraph("4.1 Criterios de aceptación", h2))
    story.append(Paragraph("Se considerará aceptable una primera fase de RH360 si quedan cubiertas las funciones esenciales de agenda, diarios, comunicaciones y administración básica de empleados. En esta primera fase no será obligatorio implantar la gestión completa de excedencias, aunque deberá quedar definida para una fase posterior.", normal))
    story.append(list_items([
        "La aplicación permite el acceso de empleados previamente dados de alta por Recursos Humanos.",
        "Recursos Humanos puede crear, modificar y dar de baja empleados en la aplicación.",
        "Los empleados pueden consultar una agenda o listado de contactos internos.",
        "Los empleados pueden buscar compañeros por nombre, departamento, área o turno.",
        "Los empleados pueden crear diarios por temática.",
        "Los empleados pueden registrar notas, incidencias u observaciones laborales.",
        "Los empleados pueden enviar comunicaciones escritas a otros compañeros.",
        "Los empleados pueden consultar las comunicaciones recibidas.",
        "El sistema mantiene un registro histórico básico de notas y comunicaciones.",
        "La gestión de excedencias queda definida funcionalmente para una fase posterior, aunque no sea obligatoria en la primera entrega del sistema.",
    ], normal))

    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def build_kanban_markdown(path):
    lines = [
        "# RH360 - Historias de usuario y tablero Kanban",
        "",
        "## Datos del proyecto",
        "",
        "- Nombre de la aplicación: RH360",
        "- Código de la aplicación: RH3",
        "- Sistema: RRHH. Recursos Humanos",
        "- Subsistema: GPRH. Gestión de Personal",
        "- Nombre del tablero Kanban: RH360",
        "",
        "## Estados obligatorios",
        "",
        "- Backlog del producto",
        "- Pendiente de desarrollo",
        "- En proceso de desarrollo",
        "- En fase de pruebas y estrés",
        "- Pendiente de despliegue",
        "- Finalizada / Histórico",
        "",
        "## Distribución propuesta de historias",
        "",
    ]
    for estado, historias in HISTORIAS.items():
        lines.append(f"### {estado}")
        lines.append("")
        for historia in historias:
            lines.append(f"- {historia}")
        lines.append("")

    lines.extend([
        "## Instrucciones para crear el tablero manualmente en GitHub",
        "",
        "1. Acceder a GitHub con la cuenta que se utilizará para la entrega.",
        "2. Entrar en el perfil de usuario y seleccionar Projects.",
        "3. Crear un nuevo proyecto con nombre RH360.",
        "4. Seleccionar vista de tipo Board o Kanban.",
        "5. Configurar la visibilidad como pública, o garantizar permiso de lectura al profesor.",
        "6. Editar el campo Status para que contenga exactamente estos valores: Backlog del producto, Pendiente de desarrollo, En proceso de desarrollo, En fase de pruebas y estrés, Pendiente de despliegue, Finalizada / Histórico.",
        "7. Crear las historias indicadas en este documento y asignar cada una al estado correspondiente.",
        "8. Copiar la URL pública del proyecto y añadirla en la entrega del aula virtual.",
        "",
        "## Comandos orientativos con GitHub CLI",
        "",
        "Estos comandos requieren tener instalado GitHub CLI (`gh`) y haber iniciado sesión con una cuenta con permisos para crear repositorios y proyectos.",
        "",
        "```powershell",
        "gh auth login",
        "gh repo create RH360 --public --description \"Aplicación corporativa para la gestión de comunicaciones internas, diarios, agenda de empleados y trámites del empleado.\" --add-readme",
        "gh project create --owner \"@me\" --title \"RH360\"",
        "```",
        "",
        "Después de crear el proyecto, se deben configurar manualmente los estados del campo Status si la versión instalada de GitHub CLI no permite modificar opciones de campos de Projects v2 directamente.",
        "",
        "Referencia oficial consultable: https://docs.github.com/issues/planning-and-tracking-with-projects/automating-your-project/using-the-api-to-manage-projects",
        "",
    ])
    path.write_text("\n".join(lines), encoding="utf-8")


def build_readme(path):
    path.write_text(
        "\n".join(
            [
                "# RH360",
                "",
                "- Nombre del proyecto: RH360",
                "- Sistema: RRHH. Recursos Humanos",
                "- Subsistema: GPRH. Gestión de Personal",
                "",
                "## Descripción",
                "",
                "Aplicación corporativa para la gestión de comunicaciones internas, diarios, agenda de empleados y trámites del empleado.",
                "",
                "## Tablero Kanban",
                "",
                "Pendiente de incorporar la URL pública del tablero Kanban de GitHub cuando esté creado.",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    docx_path = BASE_DIR / FICHERO_DOCX
    pdf_path = BASE_DIR / FICHERO_PDF
    build_docx(docx_path)
    build_fallback_pdf(pdf_path)
    build_kanban_markdown(BASE_DIR / "RH360_historias_usuario_kanban.md")
    build_readme(BASE_DIR / "README_RH360.md")
    print(docx_path)
    print(pdf_path)
